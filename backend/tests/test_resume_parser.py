from io import BytesIO

import pytest
from docx import Document

from app.services.resume_parser import ResumeParseError, extract_resume_text


def test_extracts_docx_text():
    document = Document()
    document.add_paragraph("Poojan Desai")
    document.add_paragraph("Python and TypeScript")
    buffer = BytesIO()
    document.save(buffer)

    text = extract_resume_text("resume.docx", buffer.getvalue())

    assert "Poojan Desai" in text
    assert "Python and TypeScript" in text


def test_rejects_unsupported_extension():
    with pytest.raises(ResumeParseError, match="Unsupported file type"):
        extract_resume_text("resume.txt", b"not a supported document")
